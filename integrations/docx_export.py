"""Chuyển nội dung markdown (do agent sinh) → bytes file .docx.

Agent (LLM) chỉ sinh được *text*; file .docx là gói ZIP/XML nhị phân nên phải dựng
bằng code. Module này nhận markdown đơn giản và build .docx bằng `python-docx`.

Cú pháp markdown hỗ trợ (đủ cho báo cáo PM):
- `# / ## / ###`           → heading 1/2/3
- `- ` hoặc `* `           → bullet list
- `1. ` `2. ` ...          → numbered list
- bảng kiểu `| a | b |`    → bảng docx (dòng `|---|` là separator, bỏ qua)
- `**đậm**`                → in đậm trong dòng
- dòng trống               → bỏ qua
- còn lại                  → đoạn văn thường

python-docx được import lazy: môi trường không cài lib (local/test offline) vẫn import
được module này; chỉ khi *gọi* markdown_to_docx mới cần lib. Caller nên bắt ImportError
để fallback (vd xuất .txt).
"""
from __future__ import annotations

import re

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def _is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.count("|") >= 2


def _is_separator_row(line: str) -> bool:
    # | --- | :--: | ...  (chỉ gồm | - : và khoảng trắng)
    return _is_table_row(line) and set(line.strip()) <= set("|-: ")


def _split_cells(line: str) -> list[str]:
    cells = line.strip().split("|")
    # bỏ phần tử rỗng ở 2 đầu do | mở/đóng
    if cells and cells[0].strip() == "":
        cells = cells[1:]
    if cells and cells[-1].strip() == "":
        cells = cells[:-1]
    return [c.strip() for c in cells]


def _add_runs(paragraph, text: str) -> None:
    """Thêm text vào paragraph, render **...** thành run in đậm."""
    pos = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        paragraph.add_run(m.group(1)).bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def markdown_to_docx(content: str) -> bytes:
    """Build file .docx từ markdown, trả về bytes. Cần `python-docx` (import lazy)."""
    from docx import Document  # lazy: chỉ cần khi thực sự xuất docx
    from io import BytesIO

    doc = Document()
    lines = (content or "").replace("\r\n", "\n").split("\n")
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # --- bảng markdown: gom các dòng liên tiếp bắt đầu bằng | ---
        if _is_table_row(line):
            block = []
            while i < n and _is_table_row(lines[i]):
                block.append(lines[i])
                i += 1
            rows = [_split_cells(b) for b in block if not _is_separator_row(b)]
            if rows:
                ncols = max(len(r) for r in rows)
                table = doc.add_table(rows=0, cols=ncols)
                table.style = "Light Grid Accent 1"
                for r_idx, row in enumerate(rows):
                    cells = table.add_row().cells
                    for c_idx in range(ncols):
                        txt = row[c_idx] if c_idx < len(row) else ""
                        cell_par = cells[c_idx].paragraphs[0]
                        _add_runs(cell_par, txt)
                        if r_idx == 0:  # hàng đầu = header → in đậm
                            for run in cell_par.runs:
                                run.bold = True
            continue

        # --- heading ---
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            doc.add_heading(text, level=min(max(level, 1), 4))
            i += 1
            continue

        # --- bullet ---
        if stripped.startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, stripped[2:].strip())
            i += 1
            continue

        # --- numbered ---
        m = re.match(r"^\d+\.\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, m.group(1))
            i += 1
            continue

        # --- đoạn văn thường ---
        p = doc.add_paragraph()
        _add_runs(p, stripped)
        i += 1

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
